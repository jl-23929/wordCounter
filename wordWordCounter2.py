import os
import glob
import re
import logging
from docx import Document
import tkinter
from tkinter import filedialog
import customtkinter
from PIL import Image, ImageTk
from pygame import mixer
import time
import win32com.client
import shutil

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def batch_find_replace_delete_and_remove_chars(folder_path, find_chars, replace_text, delete_chars):
    # Get all docx files in the specified folder
    docx_files = glob.glob(os.path.join(folder_path, '*.docx'))
    
    global num_files
    num_files = 0

    fileProgress = 0
    # Create the output folder     
    global modifiedFolder
    modifiedFolder = os.path.join(folder_path, 'Modified Folder')
    os.makedirs(modifiedFolder, exist_ok=True)

    files = os.listdir(folder_path)
    num_files = len(files)
    print(num_files)
    
    for docx_file in docx_files:
        try:
            logging.info(f"Processing file: {docx_file}")
            # Open each docx file
            doc = Document(docx_file)
            # Remove Bibliography
            removeReferences(doc.paragraphs)
            removeBibliography(doc.paragraphs)
            removeInTextCitations(doc.paragraphs)

            # Process paragraphs
            process_paragraphs(doc.paragraphs, find_chars, replace_text, delete_chars)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_paragraphs(cell.paragraphs, find_chars, replace_text, delete_chars)
            
            # Process headers and footers
            for section in doc.sections:
                process_paragraphs(section.header.paragraphs, find_chars, replace_text, delete_chars)
                process_paragraphs(section.footer.paragraphs, find_chars, replace_text, delete_chars)
            
            # Save the modified document with word count prepended to the file name
            global output_file_path
            output_file_name = os.path.basename(docx_file)
            output_file_path = os.path.join(modifiedFolder, output_file_name)
            doc.save(output_file_path)
            logging.info(f"Modified {docx_file} and saved to {output_file_path}")
    
            fileProgress += 1
            print(fileProgress)
        except Exception as e:
            logging.error(f"Failed to process file {docx_file}: {e}")

def process_paragraphs(paragraphs, find_chars, replace_text, delete_chars):
    for paragraph in paragraphs:
        try:
            # Replace specified characters with spaces
            updated_text = re.sub(f"[{re.escape(''.join(find_chars))}]", replace_text, paragraph.text)
            paragraph.text = updated_text
            
            # Delete specified characters only if surrounded by spaces
            updated_text = re.sub(r'\b(?:{})\b'.format('|'.join(re.escape(c) for c in delete_chars)), '', paragraph.text)
            paragraph.text = updated_text
            
            # Remove all instances of "-", "_", "–", "⇌", and "⟶"
            paragraph.text = paragraph.text.replace("-", "").replace("_", "").replace("–", "").replace("⇌", "").replace("⟶", "")
        except Exception as e:
            logging.error(f"Error processing paragraph: {e}")

def count_words_in_docx(input_folder):
    # Initialize Word application
    word_app = win32com.client.Dispatch("Word.Application")
    word_app.Visible = False  # Hide Word application window

    # Get all docx files in the specified folder
    docx_files = [file for file in os.listdir(input_folder) if file.endswith('.docx')]

    for docx_file in docx_files:
        try:
            # Open the Word document
            doc_path = os.path.join(input_folder, docx_file)
            doc = word_app.Documents.Open(doc_path)

            # Count the words in the document
            word_count = doc.ComputeStatistics(0)  # 0 for wdStatisticWords

            # Close the document without saving changes
            doc.Close(SaveChanges=False)

            # Create a copy of the document with word count appended to filename
            new_file_name = f"{word_count}_{docx_file}"
            new_file_path = os.path.join(input_folder, new_file_name)
            shutil.copyfile(doc_path, new_file_path)

            print(f"Word count for '{docx_file}': {word_count}. Copied to '{new_file_path}'")

        except Exception as e:
            print(f"Error processing file '{docx_file}': {e}")

    # Quit Word application
    word_app.Quit()

def removeBibliography(paragraphs):
    inBibliography = False

    for paragraph in paragraphs:
        if inBibliography:
            paragraph.clear()
        elif re.match(r'Bibliography', paragraph.text) or re.match(r'Reference List', paragraph.text) or re.match(r'References', paragraph.text) or re.match(r'Citations', paragraph.text) or re.match(r'References Cited', paragraph.text):
            inBibliography = True

def removeInTextCitations(paragraphs):
    for paragraph in paragraphs:
        # Searches and replaces all instances of '(text, 1111)' with ''
        updated_text = re.sub(r'\([^,]+,\s\d{4}\)', '', paragraph.text)
        paragraph.text = updated_text
        # Searches and replaces all instances of '(text, n.d.)' with ''
        updated_text = re.sub(r'\([^,]+,\sn\.d\.\)', '', paragraph.text)
        paragraph.text = updated_text
        # Searches and replaces all instances of '(text, n.d)' with ''
        updated_text = re.sub(r'\([^,]+,\sn\.d\)', '', paragraph.text)
        paragraph.text = updated_text

def removeReferences(paragraphs):
    for paragraph in paragraphs:
        # Searches for everything before references of the format '. (1111). .' and everything after.
        updated_text = re.sub(r'.*\.\s\(\d{4}\)\.\s.*', '', paragraph.text)
        paragraph.text = updated_text
        # Searches for references of the format '. (1111, Month, Day). .'
        updated_text = re.sub(r'.*\.\s\(\d{4}[^)]*\)\.\s.*', '', paragraph.text)
        paragraph.text = updated_text

def select_folder():
    global input_folder
    folder_selected = filedialog.askdirectory()
    input_folder = folder_selected

# Define characters to find and replace with space (excluding "-", "_", "–", "⇌", and "⟶")
find_chars = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '0', ',', '.', '?', '!', ':', ';', '(', ')', '[', ']', '{', '}', '/', '\\', '*', '+', '=', '|', '&', '^', '%', '@', '~', '`', '"', "'", '°', '𝜃', '−', '×', '±', '≈', '∆', '>', '<', '>=', '<=', '=']

replace_text = ' '

# Define characters to delete (only if surrounded by spaces)
delete_chars = ['M', 'V', 'Z', 'C', 'Q', 'Cu', 'Zn', 'Ag', 'NO', 'KNO', 'MnO', 'NaCl', 'kPa', 'mL', 'L', 'aq', 'l', 's', 'g', 'x', 'J' 'KWh' 'kWh' 'cm' 'm', 'kW', 'W', 'MW', 'RPM', 'rpm']

def validate_and_get_word_limit(entry_widget):
    try:
        word_limit = int(entry_widget.get())
        if word_limit <= 0:
            entry_widget.configure(placeholder_text_color='red')
            raise ValueError("Word limit must be greater than 0.")
        entry_widget.configure(placeholder_text_color='black')
        return word_limit
    except ValueError:
        entry_widget.configure(placeholder_text_color='red')
        return None

def start_sorting():
    word_limit = validate_and_get_word_limit(wordLimitEntry)
    if word_limit is None:
        return
    #updateGif()
    batch_find_replace_delete_and_remove_chars(input_folder, find_chars, replace_text, delete_chars)
    count_words_in_docx(modifiedFolder)
    os.startfile(output_file_path)
    playCount()

window = customtkinter.CTk()

def get_absolute_path(relative_path):
    base_path = os.path.dirname(__file__)
    return os.path.join(base_path, relative_path)

window.geometry("700x400+600+300")
window.title("Word Document Modifier")
#window.iconbitmap(get_absolute_path(".ico"))

fontStyle = customtkinter.CTkFont(family="Arial", size=20)

instructionsLabel = customtkinter.CTkLabel(window, text="This program will remove all in-text citations, the Bibliography, and other references.\nIt will also replace unwanted characters and count the words in each document.", font=fontStyle)
instructionsLabel.pack(pady=10)

fontStyle = customtkinter.CTkFont(family="Arial", size=12)

wordLimitLabel = customtkinter.CTkLabel(window, text="Word Limit:", font=fontStyle)
wordLimitLabel.pack(pady=5)
wordLimitEntry = customtkinter.CTkEntry(window, placeholder_text="Enter word limit", font=fontStyle)
wordLimitEntry.pack(pady=5)

selectFolderButton = customtkinter.CTkButton(window, text="Select Folder", command=select_folder, font=fontStyle)
selectFolderButton.pack(pady=10)

startButton = customtkinter.CTkButton(window, text="Start Sorting", command=start_sorting, font=fontStyle)
startButton.pack(pady=10)
mixer.init()

def playCount():
    sound = get_absolute_path("sounds\\count.mp3")
    mixer.music.load(sound)
    mixer.music.play()

window.mainloop()
