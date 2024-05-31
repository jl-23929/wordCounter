import os
import glob
import re
import logging
from docx import Document
import tkinter
from tkinter import filedialog, messagebox
from tkinter import *
import customtkinter
from PIL import Image


# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def batch_find_replace_delete_and_remove_chars(folder_path, find_chars, replace_text, delete_chars):
    # Get all docx files in the specified folder
    docx_files = glob.glob(os.path.join(folder_path, '*.docx'))
    
    global num_files
    num_files = 0

    fileProgress = 0
    # Create the output     
    above_2000_folder = os.path.join(folder_path, 'Above ' + str(wordLimitEntry.get())) 
    already_under_2000_folder = os.path.join(folder_path, 'Already under ' + str(wordLimitEntry.get()))
    os.makedirs(above_2000_folder, exist_ok=True)
    os.makedirs(already_under_2000_folder, exist_ok=True)

    files = os.listdir(folder_path)
    num_files = len(files)
    print(num_files)
    
    for docx_file in docx_files:
        try:
            logging.info(f"Processing file: {docx_file}")
            # Open each docx file
            doc = Document(docx_file)
            # Remove Bibliography
            removeReferences(doc.paragraphs)

            removeBibliography(doc.paragraphs)

            removeInTextCitations(doc.paragraphs)

            # Process paragraphs
            process_paragraphs(doc.paragraphs, find_chars, replace_text, delete_chars)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_paragraphs(cell.paragraphs, find_chars, replace_text, delete_chars)
            
            # Process headers and footers
            for section in doc.sections:
                process_paragraphs(section.header.paragraphs, find_chars, replace_text, delete_chars)
                process_paragraphs(section.footer.paragraphs, find_chars, replace_text, delete_chars)
            
            # Word count
            word_count = get_word_count(doc)
            
            # Determine the output subfolder based on word count
            if word_count > int(wordLimitEntry.get()):
                output_subfolder = above_2000_folder
            else:
                output_subfolder = already_under_2000_folder
            
            # Save the modified document with word count prepended to the file name
            output_file_name = f"{word_count}_{os.path.basename(docx_file)}"
            output_file_path = os.path.join(output_subfolder, output_file_name)
            doc.save(output_file_path)
            logging.info(f"Modified {docx_file} and saved to {output_file_path}")
        
            fileProgress = fileProgress+1
            #progressBar.step()

        except Exception as e:
            logging.error(f"Failed to process file {docx_file}: {e}")

def process_paragraphs(paragraphs, find_chars, replace_text, delete_chars):
    for paragraph in paragraphs:
        try:
            # Replace specified characters with spaces
            updated_text = re.sub(f"[{re.escape(''.join(find_chars))}]", replace_text, paragraph.text)
            paragraph.text = updated_text
            
            # Delete specified characters only if surrounded by spaces
            updated_text = re.sub(r'\b(?:{})\b'.format('|'.join(re.escape(c) for c in delete_chars)), '', paragraph.text)
            paragraph.text = updated_text
            
            # Remove all instances of "-", "_", "–", "⇌", and "⟶"
            paragraph.text = paragraph.text.replace("-", "").replace("_", "").replace("–", "").replace("⇌", "").replace("⟶", "")
        except Exception as e:
            logging.error(f"Error processing paragraph: {e}")

def get_word_count(doc):
    try:
        word_count = 0
        for paragraph in doc.paragraphs:
            word_count += len(paragraph.text.split())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        word_count += len(paragraph.text.split())
        return word_count
    except Exception as e:
        logging.error(f"Error counting words: {e}")
        return 0
    
input_folder = ''

def removeBibliography(paragraphs):
    inBibliography = False

    for paragraph in paragraphs:

        if inBibliography:
            paragraph.clear()
        elif re.match(r'Bibliography', paragraph.text) or re.match(r'Reference List', paragraph.text) or re.match(r'References', paragraph.text) or re.match(r'Citations', paragraph.text) or re.match(r'References Cited', paragraph.text):
            
            inBibliography = True

def removeInTextCitations(paragraphs):

    for paragraph in paragraphs:
        
        # Searches and replaces all instances of '(text, 1111)' with ''
        updated_text = re.sub(r'\([^,]+,\s\d{4}\)', '', paragraph.text)
        paragraph.text = updated_text


        #Searches and replaces all instances of '(text, n.d.)' with ''
        updated_text = re.sub(r'\([^,]+,\sn\.d\.\)', '', paragraph.text)
        paragraph.text = updated_text

        #Searches and replaces all instances of '(text, n.d)' with ''
        updated_text = re.sub(r'\([^,]+,\sn\.d\)', '', paragraph.text)
        paragraph.text = updated_text

def removeReferences(paragraphs):

    for paragraph in paragraphs:
        #Searches for everything before references of the format '. (1111). .' and everything after.
        updated_text = re.sub(r'.*\.\s\(\d{4}\)\.\s.*', '', paragraph.text)
        paragraph.text = updated_text
        
        #Searches for references of the format '. (1111, Month, Day). .'
        updated_text = re.sub(r'.*\.\s\(\d{4}[^)]*\)\.\s.*', '', paragraph.text)
        paragraph.text = updated_text


def select_folder():
    global input_folder
    folder_selected = filedialog.askdirectory()
    input_folder = folder_selected

# Define characters to find and replace with space (excluding "-", "_", "–", "⇌", and "⟶")
find_chars = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '0', ',', '.', '?', '!', ':', ';', '(', ')', '[', ']', '{', '}', '/', '\\', '*', '+', '=', '|', '&', '^', '%', '@', '~', '`', '"', "'", '°', '𝜃', '−', '×', '±', '≈', '∆', '>', '<', '>=', '<=', '=']

# Define text to replace found characters with
replace_text = ' '

# Define characters to delete (only if surrounded by spaces)
delete_chars = ['M', 'V', 'Z', 'C', 'Q', 'Cu', 'Zn', 'Ag', 'NO', 'KNO', 'MnO', 'NaCl', 'kPa', 'mL', 'L', 'aq', 'l', 's', 'g', 'x', 'J' 'KWh' 'kWh' 'cm' 'm', 'kW', 'W', 'MW', 'RPM', 'rpm']

def start_sorting():
    
    batch_find_replace_delete_and_remove_chars(input_folder, find_chars, replace_text, delete_chars)
    messagebox.showinfo("Success", "Files have been sorted successfully.")


window = customtkinter.CTk()

window.geometry("700x400")

window.title("Cookie Monster")
window.resizable(0,0)
customtkinter.set_default_color_theme("blue")
def selectFolder():
    global input_folder
    folder_selected = filedialog.askdirectory()
    input_folder = folder_selected

column1 = 0.1
column2 = 0.2
column3 = 0.35

image_path = r"C:\Users\james680384\Pictures\Picture1.png"

pil_image = Image.open(image_path)

image = customtkinter.CTkImage(light_image=pil_image, dark_image=pil_image, size=(200,200))
imageLabel = customtkinter.CTkLabel(window, image=image, text="")
imageLabel.place(relx = 0.25, rely = 0.7, anchor=CENTER)


bold = customtkinter.CTkFont(family="Arial Black", size=32, )

infoHeading = customtkinter.CTkLabel(window, text="Cookie Monster", font=bold, text_color="#004f98")
infoHeading.place(relx=0.5, rely=column1, anchor=CENTER)

infoLabel = customtkinter.CTkLabel(window, text="To use Cookie Monster, ..................")
infoLabel.place(relx=0.5, rely=column2, anchor=CENTER)

selectFolderLabel = customtkinter.CTkLabel(window, text="Select Folder:")
selectFolderLabel.place(relx=0.2, rely = column3, anchor=CENTER)

folderEntry = customtkinter.CTkEntry(window, placeholder_text="Enter a path or browse", width=160)
folderEntry.place(relx = 0.5, rely = column3, anchor=CENTER)

browseButton = customtkinter.CTkButton(master=window, text="Browse", command=selectFolder)
browseButton.place(relx=0.8, rely=column3, anchor=CENTER)

wordLimitEntry = customtkinter.CTkEntry(window, placeholder_text="Word Limit")
wordLimitEntry.place(relx=0.5, rely = 0.5, anchor=CENTER)


sortButton = customtkinter.CTkButton(master=window, text="Start Sorting", command=start_sorting)
sortButton.place(relx=0.5, rely=0.75, anchor = CENTER)

window.mainloop()