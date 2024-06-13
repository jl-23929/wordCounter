#WordWordCounter.py>
import os
import shutil
import win32com.client
import re
from docx import Document

def count_words_in_docx(input_folder, wordLimit):
    # Initialize Word application
    word_app = win32com.client.Dispatch("Word.Application")
    word_app.Visible = False  # Hide Word application window

    docx_files1 = [file for file in os.listdir(input_folder) if file.endswith('.docx')]

    for docx_file in docx_files1:
        try:
            doc_path = os.path.join(input_folder, docx_file)
            doc = Document(doc_path)

            #removences(doc_path)
            removeBibliography(doc_path)
        except Exception as e:
            print(f"Error processing file '{docx_file}': {e}")

    # Get all docx files in the specified folder
    docx_files = [file for file in os.listdir(input_folder) if file.endswith('.docx')]
    for docx_file in docx_files:
        try:
            # Open the Word document
            doc_path = os.path.join(input_folder, docx_file)
            doc = word_app.Documents.Open(doc_path)
            patterns = [',', '=', '𝜃', '0', '1', '2', '3', '4', '5', '6', '7', '8', '9', '.', '?', ':', ';', '(', ')', '[', ']', '{', '}', '/', '*', '+', '=', '|', '&', '%', '@',
                         '~', '`', '"', '°', '−', '×', '±', '≈', '∆', '>', '<', '>=', '<=', '=', 'ϕ', 'φ', 'Φ', 'Ω', 'Å', '𝜙', ' NaCl ', ' kPa ', ' mL ', ' L ', ' aq ', ' l ', ' s ', 
                         ' g ', ' x ', ' KWh ', ' kWh ', ' cm ', ' m ', ' kW ', ' W ', ' MW ', ' RPM ', ' rpm ', ' CO2 ', " ' ", ' MHz ', ' km ', ' nm ', ' mV ', ' THz ', ' eV ', 
                         ' keV ', ' MeV ', ' J ', ' Hz ', ' kHz ']
            
            
            #Had to remove !, \\  and ^

            content = doc.Content

            for pattern in patterns:
                find = content.Find
                find.ClearFormatting()

                find.Text = pattern
                find.Replacement.ClearFormatting()
                find.Replacement.Text = ""
                find.Execute(Replace=2, MatchWildcards=False)
                print(f"Removed {pattern} from {docx_file}")

                    #If not work: try story.Delete()
            
            find.Text = "-"
            find.ClearFormatting()
            find.Replacement.ClearFormatting()
            find.Replacement.Text = " "
            find.Execute(Replace=2, MatchWildcards=False)
            print(f"Removed - from {docx_file}")

            doc.SaveAs(f"Processed_{docx_file}")

                
            # Count the words in the document
            word_count = doc.ComputeStatistics(0, True)  # 0 for wdStatisticWords, True for including textboxes

            # Close the document without saving changes
            doc.Close(SaveChanges=False)

            # Create a copy of the document with word count appended to filename

            #above_2000_folder = os.path.join(input_folder, 'Above ' + str(wordLimit)) 
            above_2000_folder = os.path.join(input_folder, "Processed")
            #already_under_2000_folder = os.path.join(input_folder, 'Already under ' + str(wordLimit))
            os.makedirs(above_2000_folder, exist_ok=True)
            #os.makedirs(already_under_2000_folder, exist_ok=True)
            
            if word_count > int(wordLimit):
                output_subfolder = above_2000_folder
            else:
                #output_subfolder = already_under_2000_folder
                output_subfolder = above_2000_folder
            
            new_file_name = f"{word_count}_{docx_file}"
            new_file_path = os.path.join(output_subfolder, new_file_name)

            shutil.copyfile(doc_path, new_file_path)

            print(f"Word count for '{docx_file}': {word_count}. Copied to '{new_file_path}'")

        except Exception as e:
            print(f"Error processing file '{docx_file}': {e}")

    # Quit Word application
    word_app.Quit()
    
def removeReferences(file_path):
    doc = Document(file_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run_text = run.text
            # Search for references and remove them while preserving formatting
            new_text = re.sub(r'.*\.\s\(\d{4}\)\.\s.*', '', run_text)
            new_text = re.sub(r'.*\.\s\(\d{4}[^)]*\)\.\s.*', '', new_text)
            new_text = re.sub(r'.*\.\s\(n\.d\.\)\.\s.*', '', new_text)
            if new_text != run_text:
                run.text = new_text

    # Save changes back to the document
    doc.save(file_path)
def removeBibliography(file_path):

    doc = Document(file_path)

    inBibliography = False

    paragraphs = doc.paragraphs

    for paragraph in paragraphs:

        if inBibliography:
            paragraph.clear()
        elif re.match(r'Bibliography', paragraph.text) or re.match(r'Reference List', paragraph.text) or re.match(r'References', paragraph.text) or re.match(r'Citations', paragraph.text) or re.match(r'References Cited', paragraph.text):
            
            inBibliography = True
    doc.save(file_path)