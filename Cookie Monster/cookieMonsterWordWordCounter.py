import os
import glob
import re
import logging
from docx import Document
import tkinter
from tkinter import filedialog, messagebox
from tkinter import *
import customtkinter
from PIL import Image, ImageTk
from pygame import mixer
import time
import win32com.client
import shutil
import WordWordCounter
# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def batch_find_replace_delete_and_remove_chars(folder_path, find_chars, replace_text, delete_chars):
    # Get all docx files in the specified folder
    docx_files = glob.glob(os.path.join(folder_path, '*.docx'))
    
    global num_files
    num_files = 0

    fileProgress = 0
    # Create the output     
    global modifiedFolder
    modifiedFolder = os.path.join(folder_path, 'Modified Folder')
    os.makedirs(modifiedFolder, exist_ok=True)

    files = os.listdir(folder_path)
    num_files = len(files)
    print(num_files)
    
    for docx_file in docx_files:
        try:
            logging.info(f"Processing file: {docx_file}")
            # Open each docx file
            doc = Document(docx_file)
            # Remove Bibliography
            removeReferences(doc.paragraphs)

            removeBibliography(doc.paragraphs)

            removeInTextCitations(doc.paragraphs)

            # Process paragraphs
            process_paragraphs(doc.paragraphs, find_chars, replace_text, delete_chars)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_paragraphs(cell.paragraphs, find_chars, replace_text, delete_chars)
            
            # Process headers and footers
            for section in doc.sections:
                process_paragraphs(section.header.paragraphs, find_chars, replace_text, delete_chars)
                process_paragraphs(section.footer.paragraphs, find_chars, replace_text, delete_chars)
            # Save the modified document with word count prepended to the file name
            global output_file_path
            output_file_name = os.path.basename(docx_file)
            output_file_path = os.path.join(modifiedFolder, output_file_name)
            doc.save(output_file_path)
            logging.info(f"Modified {docx_file} and saved to {output_file_path}")
    
            fileProgress += 1
            print(fileProgress)
        except Exception as e:
            logging.error(f"Failed to process file {docx_file}: {e}")
def process_paragraphs(paragraphs, find_chars, replace_text, delete_chars):
    for paragraph in paragraphs:
        try:
            # Replace specified characters with spaces
            updated_text = re.sub(f"[{re.escape(''.join(find_chars))}]", replace_text, paragraph.text)
            paragraph.text = updated_text
            
            # Delete specified characters only if surrounded by spaces
            updated_text = re.sub(r'\b(?:{})\b'.format('|'.join(re.escape(c) for c in delete_chars)), '', paragraph.text)
            paragraph.text = updated_text
            
            # Remove all instances of "-", "_", "–", "⇌", and "⟶"
            paragraph.text = paragraph.text.replace("-", "").replace("_", "").replace("–", "").replace("⇌", "").replace("⟶", "")
        except Exception as e:
            logging.error(f"Error processing paragraph: {e}")




def count_words_in_docx(input):
    # Initialize Word application
    word_app = win32com.client.Dispatch("Word.Application")
    word_app.Visible = False  # Hide Word application window

    # Get all docx files in the specified folder
    docx_files = [file for file in os.listdir(input) if file.endswith('.docx')]

    for docx_file in docx_files:
        try:
            # Open the Word document
            doc_path = os.path.join(input, docx_file)
            doc = word_app.Documents.Open(doc_path)

            # Count the words in the document
            word_count = doc.ComputeStatistics(0)  # 0 for wdStatisticWords

            # Close the document without saving changes
            doc.Close(SaveChanges=False)

            # Create a copy of the document with word count appended to filename
            new_file_name = f"{word_count}_{docx_file}"
            new_file_path = os.path.join(input, new_file_name)
            shutil.copyfile(doc_path, new_file_path)

            print(f"Word count for '{docx_file}': {word_count}. Copied to '{new_file_path}'")

        except Exception as e:
            print(f"Error processing file '{docx_file}': {e}")

    # Quit Word application
    word_app.Quit()



input_folder = ''

def removeBibliography(paragraphs):
    inBibliography = False

    for paragraph in paragraphs:

        if inBibliography:
            paragraph.clear()
        elif re.match(r'Bibliography', paragraph.text) or re.match(r'Reference List', paragraph.text) or re.match(r'References', paragraph.text) or re.match(r'Citations', paragraph.text) or re.match(r'References Cited', paragraph.text):
            
            inBibliography = True

def removeInTextCitations(paragraphs):

    for paragraph in paragraphs:
        
        # Searches and replaces all instances of '(text, 1111)' with ''
        updated_text = re.sub(r'\([^,]+,\s\d{4}\)', '', paragraph.text)
        paragraph.text = updated_text


        #Searches and replaces all instances of '(text, n.d.)' with ''
        updated_text = re.sub(r'\([^,]+,\sn\.d\.\)', '', paragraph.text)
        paragraph.text = updated_text

        #Searches and replaces all instances of '(text, n.d)' with ''
        updated_text = re.sub(r'\([^,]+,\sn\.d\)', '', paragraph.text)
        paragraph.text = updated_text

def removeReferences(paragraphs):

    for paragraph in paragraphs:
        #Searches for everything before references of the format '. (1111). .' and everything after.
        updated_text = re.sub(r'.*\.\s\(\d{4}\)\.\s.*', '', paragraph.text)
        paragraph.text = updated_text
        
        #Searches for references of the format '. (1111, Month, Day). .'
        updated_text = re.sub(r'.*\.\s\(\d{4}[^)]*\)\.\s.*', '', paragraph.text)
        paragraph.text = updated_text


def select_folder():
    global input_folder
    folder_selected = filedialog.askdirectory()
    input_folder = folder_selected

# Define characters to find and replace with space (excluding "-", "_", "–", "⇌", and "⟶")
find_chars = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '0', ',', '.', '?', '!', ':', ';', '(', ')', '[', ']', '{', '}', '/', '\\', '*', '+', '=', '|', '&', '^', '%', '@', '~', '`', '"', "'", '°', '𝜃', '−', '×', '±', '≈', '∆', '>', '<', '>=', '<=', '=']

replace_text = ' '

# Define characters to delete (only if surrounded by spaces)
delete_chars = ['M', 'V', 'Z', 'C', 'Q', 'Cu', 'Zn', 'Ag', 'NO', 'KNO', 'MnO', 'NaCl', 'kPa', 'mL', 'L', 'aq', 'l', 's', 'g', 'x', 'J' 'KWh' 'kWh' 'cm' 'm', 'kW', 'W', 'MW', 'RPM', 'rpm']

def validate_and_get_word_limit(entry_widget):
    try:
        word_limit = int(entry_widget.get())
        if word_limit <= 0:
            entry_widget.configure(placeholder_text_color='red')
            raise ValueError("Word limit must be greater than 0.")
        entry_widget.configure(placeholder_text_color='black')
        return word_limit
    except ValueError:
        entry_widget.configure(placeholder_text_color='red')
        return None

def start_sorting():
    word_limit = validate_and_get_word_limit(wordLimitEntry)
    if word_limit is None:
        return
    #updateGif()
    batch_find_replace_delete_and_remove_chars(input_folder, find_chars, replace_text, delete_chars)
    count_words_in_docx(modifiedFolder)
    os.startfile(output_file_path)
    playCount()


window = customtkinter.CTk()

def get_absolute_path(relative_path):
    base_path = os.path.dirname(__file__)
    return os.path.join(base_path, relative_path)

window.geometry("700x400+600+300")
icon_path = get_absolute_path("Data/Monster.ico")
window.iconbitmap(icon_path)
window.title("Cookie Monster")
window.resizable(0,0)
customtkinter.set_default_color_theme("blue")
def selectFolder():
    global input_folder
    folder_selected = filedialog.askdirectory()
    input_folder = folder_selected

def get_asset_path(asset_type, filename):
    base_path = os.path.dirname(__file__)
    return os.path.join(base_path, 'assets', asset_type, filename)

column1 = 0.1
column2 = 0.3
column3 = 0.45

image_path = get_absolute_path("Data/Cookie Monster Image.png")
soundIconPath = get_absolute_path("Data/noun-play-button-6441783-FFFFFF.png")
soundStopPath = get_absolute_path("Data/noun-stop-button-4906815-FFFFFF.png")

pil_soundStopPath = Image.open(soundStopPath)


pil_soundIconPath = Image.open(soundIconPath)

pil_image = Image.open(image_path)

def playIntro():
    mixer.music.load(r"Data\Cookie Instructions.mp3")
    mixer.music.play()

def stopIntro():
    mixer.music.stop()

def playCount(): 
    mixer.music.load(get_absolute_path("Data\Documents Completed-[AudioTrimmer.com]-[AudioTrimmer.com].mp3"))
    mixer.music.play()
    while mixer.music.get_busy():
        time.sleep(1)
    mixer.music.load(get_absolute_path("Data\Count's Laugh 1.mp3"))
    mixer.music.play()
mixer.init()


stopImage = customtkinter.CTkImage(light_image=pil_soundStopPath, dark_image=pil_soundStopPath, size=(50,50))
soundImage = customtkinter.CTkImage(light_image=pil_soundIconPath, dark_image=pil_soundIconPath, size=(50, 50))
image = customtkinter.CTkImage(light_image=pil_image, dark_image=pil_image, size=(200,200))
imageLabel = customtkinter.CTkLabel(window, image=image, text="")
imageLabel.place(relx = 0.2, rely = 0.75, anchor=CENTER)

soundImageButton = customtkinter.CTkButton(window, image=soundImage, text="", width=55, command=playIntro)

soundImageButton.place(relx = 0.8, rely=0.3, anchor=CENTER)
soundStopButton = customtkinter.CTkButton(window, image=stopImage, text="", width=55, command=stopIntro)
soundStopButton.place(relx=0.9, rely=0.3, anchor=CENTER)
bold = customtkinter.CTkFont(family="Arial Black", size=32)
body = customtkinter.CTkFont(family="Arial", size=16)
boldBody = customtkinter.CTkFont(family="Arial", size=25, weight="bold")


infoHeading = customtkinter.CTkLabel(window, text="Cookie Monster", font=bold, text_color="#004f98")
infoHeading.place(relx=0.5, rely=column1, anchor=CENTER)

infoLabel = customtkinter.CTkLabel(window, text="Select a folder of Word documents to \n process and enter a word limit. For \n more instructions press the play button.", font=body)
infoLabel.place(relx=0.5, rely=column2, anchor=CENTER)

selectFolderLabel = customtkinter.CTkLabel(window, text="Select Folder:", font=body)
selectFolderLabel.place(relx=0.2, rely = column3, anchor=CENTER)

folderEntry = customtkinter.CTkEntry(window, placeholder_text="Enter a path or browse", width=270, font=body)
folderEntry.place(relx = 0.5, rely = column3, anchor=CENTER)

browseButton = customtkinter.CTkButton(master=window, text="Browse", command=selectFolder, font=body)
browseButton.place(relx=0.85, rely=column3, anchor=CENTER)

wordLimitEntry = customtkinter.CTkEntry(window, placeholder_text='Enter Word Limit', font=body)
wordLimitEntry.place(relx=0.5, rely = 0.6, anchor=CENTER)

sortButton = customtkinter.CTkButton(master=window, text="Cookie Time!", command=start_sorting, font=boldBody)
sortButton.place(relx=0.5, rely=0.85, anchor = CENTER)

window.mainloop()