import os
import glob
import re
import logging
from docx import Document

class cookieMonsterBack():

    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

    def batch_find_replace_delete_and_remove_chars(folder_path, find_chars, replace_text, delete_chars):
    # Get all docx files in the specified folder
        docx_files = glob.glob(os.path.join(folder_path, '*.docx'))
    
        num_files = 0

        progress = 0
        # Create the output     
        above_2000_folder = os.path.join(folder_path, 'Above ' + str(word_limit_entry.get())) 
        already_under_2000_folder = os.path.join(folder_path, 'Already under ' + str(word_limit_entry.get()))
        os.makedirs(above_2000_folder, exist_ok=True)
        os.makedirs(already_under_2000_folder, exist_ok=True)

        for docx_file in docx_files:
        
            num_files = num_files+1
    
        for docx_file in docx_files:
            try:
                logging.info(f"Processing file: {docx_file}")
                # Open each docx file
                doc = Document(docx_file)
                
                # Process paragraphs
                process_paragraphs(doc.paragraphs, find_chars, replace_text, delete_chars)
                
                # Process tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            process_paragraphs(cell.paragraphs, find_chars, replace_text, delete_chars)
                
                # Process headers and footers
                for section in doc.sections:
                    process_paragraphs(section.header.paragraphs, find_chars, replace_text, delete_chars)
                    process_paragraphs(section.footer.paragraphs, find_chars, replace_text, delete_chars)
                
                # Word count
                word_count = get_word_count(doc)
                
                # Determine the output subfolder based on word count
                if word_count > int(word_limit_entry.get()):
                    output_subfolder = above_2000_folder
                else:
                    output_subfolder = already_under_2000_folder
                
                # Save the modified document with word count prepended to the file name
                output_file_name = f"{word_count}_{os.path.basename(docx_file)}"
                output_file_path = os.path.join(output_subfolder, output_file_name)
                doc.save(output_file_path)
                logging.info(f"Modified {docx_file} and saved to {output_file_path}")
            
                progress+1

                

            except Exception as e:
                logging.error(f"Failed to process file {docx_file}: {e}")

def process_paragraphs(paragraphs, find_chars, replace_text, delete_chars):
    for paragraph in paragraphs:
        try:
            # Replace specified characters with spaces
            updated_text = re.sub(f"[{re.escape(''.join(find_chars))}]", replace_text, paragraph.text)
            paragraph.text = updated_text
            
            # Delete specified characters only if surrounded by spaces
            updated_text = re.sub(r'\b(?:{})\b'.format('|'.join(re.escape(c) for c in delete_chars)), '', paragraph.text)
            paragraph.text = updated_text
            
            # Remove all instances of "-", "_", "–", "⇌", and "⟶"
            paragraph.text = paragraph.text.replace("-", "").replace("_", "").replace("–", "").replace("⇌", "").replace("⟶", "")
        except Exception as e:
            logging.error(f"Error processing paragraph: {e}")

def get_word_count(doc):
    try:
        word_count = 0
        for paragraph in doc.paragraphs:
            word_count += len(paragraph.text.split())
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        word_count += len(paragraph.text.split())
        return word_count
    except Exception as e:
        logging.error(f"Error counting words: {e}")
        return 0
    
find_chars = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '0', ',', '.', '?', '!', ':', ';', '(', ')', '[', ']', '{', '}', '/', '\\', '*', '+', '=', '|', '&', '^', '%', '@', '~', '`', '"', "'", '°', '𝜃', '−', '×', '±', '≈', '∆', '>', '<', '>=', '<=', '=']

replace_text = ' '

delete_chars = ['M', 'V', 'Z', 'C', 'Q', 'Cu', 'Zn', 'Ag', 'NO', 'KNO', 'MnO', 'NaCl', 'kPa', 'mL', 'L', 'aq', 'l', 's', 'g', 'x']

# \d does any number, 